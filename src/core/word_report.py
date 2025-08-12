from typing import Dict, Any, List
from docx import Document

def build_detailed_docx(report: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading('ADGM Corporate Agent - Detailed Analysis', 0)
    p = doc.add_paragraph()
    p.add_run('Process: ').bold = True
    p.add_run(report.get('process', 'Unknown'))

    # Summary metrics
    doc.add_heading('Summary', level=1)
    for label, key in [
        ('Documents Uploaded', 'documents_uploaded'),
        ('Required Documents', 'required_documents'),
        ('Missing Documents', 'missing_documents'),
    ]:
        value = report.get(key)
        if isinstance(value, list):
            value = ', '.join(value)
        doc.add_paragraph(f"{label}: {value}")

    # Findings
    doc.add_heading('Findings', level=1)
    issues: List[Dict[str, Any]] = report.get('issues_found', [])
    if not issues:
        doc.add_paragraph('No issues detected.')
    else:
        for i, issue in enumerate(issues, 1):
            doc.add_heading(f"{i}. {issue.get('issue','')}", level=2)
            doc.add_paragraph(f"Document: {issue.get('document','N/A')}")
            doc.add_paragraph(f"Severity: {issue.get('severity','N/A')}")
            rat = issue.get('rationale') or ''
            if rat:
                doc.add_paragraph('Rationale: ' + rat)
            sug = issue.get('suggestion') or ''
            if sug:
                doc.add_paragraph('Suggestion: ' + sug)
            cits = issue.get('citations') or []
            if cits:
                doc.add_paragraph('Sources: ' + ', '.join(cits))

    # Per document
    doc.add_heading('Per Document', level=1)
    by_doc: Dict[str, List[Dict[str, Any]]] = {}
    for it in issues:
        by_doc.setdefault(it.get('document','Unknown'), []).append(it)
    for name, arr in by_doc.items():
        doc.add_heading(name, level=2)
        for j, it in enumerate(arr, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{j}. {it.get('issue','')} ({it.get('severity','')}) ")
            r.bold = True
            if it.get('suggestion'):
                doc.add_paragraph('Suggestion: ' + it['suggestion'])

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
