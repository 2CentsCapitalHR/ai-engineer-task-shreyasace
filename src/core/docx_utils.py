from typing import List, Dict, Any, Tuple
from io import BytesIO
from docx import Document
from docx.shared import RGBColor


def extract_text(doc_bytes: bytes) -> Tuple[str, Document]:
    """Return full text and the loaded Document object."""
    doc = Document(BytesIO(doc_bytes))
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts), doc


def insert_comments_and_return_bytes(doc_bytes: bytes, comments: List[Dict[str, Any]]) -> bytes:
    """
    Pseudo-inline comments: we append short red inline notes next to the first matching text per issue.
    Fallback-friendly approach using python-docx (true Word comments require low-level XML).
    Expected comment dict keys: { 'issue', 'location', 'suggestion', 'citations' }
    """
    text, doc = extract_text(doc_bytes)

    # Build a quick paragraph index
    para_texts = [p.text for p in doc.paragraphs]

    for c in comments:
        note = c.get("issue", "Issue")
        suggestion = c.get("suggestion")
        cite = ", ".join(c.get("citations", []) or [])
        anchor = c.get("location") or ""

        # Find first paragraph containing the anchor; else use first paragraph
        idx = 0
        if anchor:
            for i, t in enumerate(para_texts):
                if anchor.lower() in (t or "").lower():
                    idx = i
                    break
        if idx >= len(doc.paragraphs):
            idx = len(doc.paragraphs) - 1 if doc.paragraphs else 0

        if not doc.paragraphs:
            # Create at least one paragraph
            doc.add_paragraph("")

        p = doc.paragraphs[idx]
        run = p.add_run(f"  [Comment: {note}]")
        run.font.color.rgb = RGBColor(200, 0, 0)
        if suggestion:
            run2 = p.add_run(f" Suggestion: {suggestion}")
            run2.font.color.rgb = RGBColor(120, 0, 0)
        if cite:
            run3 = p.add_run(f" Sources: {cite}")
            run3.font.color.rgb = RGBColor(90, 90, 90)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
